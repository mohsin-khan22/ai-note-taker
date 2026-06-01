from docx import Document
from fpdf import FPDF
from io import BytesIO


def _format_entities(entities: dict | None) -> str:
    if not entities:
        return ""
    lines = []
    for person in entities.get("people", []):
        lines.append(f"- {person} (Person)")
    for org in entities.get("organizations", []):
        lines.append(f"- {org} (Organization)")
    for date in entities.get("dates", []):
        lines.append(f"- {date} (Date)")
    for loc in entities.get("locations", []):
        lines.append(f"- {loc} (Location)")
    return "\n".join(lines)


class MeetingExporter:
    @staticmethod
    def to_txt(data: dict) -> BytesIO:
        output = BytesIO()
        text = f"Meeting Title: {data['title']}\nDate: {data['date']}\n\n"
        text += "SUMMARY\n" + "=" * 10 + "\n" + data["summary"] + "\n\n"
        text += "KEY POINTS\n" + "=" * 10 + "\n"
        text += "\n".join([f"- {p}" for p in data["key_points"]]) + "\n\n"
        text += "ACTION ITEMS\n" + "=" * 10 + "\n"
        text += "\n".join([f"- {a}" for a in data["action_items"]]) + "\n\n"

        entity_text = _format_entities(data.get("entities"))
        if entity_text:
            text += "KEY PEOPLE & ORGANIZATIONS\n" + "=" * 10 + "\n" + entity_text + "\n\n"

        text += "TRANSCRIPT\n" + "=" * 10 + "\n" + data["transcript"]
        output.write(text.encode("utf-8"))
        output.seek(0)
        return output

    @staticmethod
    def to_docx(data: dict) -> BytesIO:
        doc = Document()
        doc.add_heading(data["title"], 0)
        doc.add_paragraph(f"Date: {data['date']}")

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data["summary"])

        doc.add_heading("Key Points", level=1)
        for point in data["key_points"]:
            doc.add_paragraph(point, style="List Bullet")

        doc.add_heading("Action Items", level=1)
        for item in data["action_items"]:
            doc.add_paragraph(item, style="List Bullet")

        entity_text = _format_entities(data.get("entities"))
        if entity_text:
            doc.add_heading("Key People & Organizations", level=1)
            for line in entity_text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.lstrip("- "), style="List Bullet")

        doc.add_heading("Full Transcript", level=1)
        doc.add_paragraph(data["transcript"])

        target = BytesIO()
        doc.save(target)
        target.seek(0)
        return target

    @staticmethod
    def to_pdf(data: dict) -> BytesIO:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(40, 10, data["title"][:80])
        pdf.ln(10)
        pdf.set_font("Arial", size=10)
        pdf.cell(40, 10, f"Date: {data['date']}")
        pdf.ln(20)

        sections = [
            ("Summary", data["summary"]),
            ("Key Points", "\n".join([f"- {p}" for p in data["key_points"]])),
            ("Action Items", "\n".join([f"- {a}" for a in data["action_items"]])),
        ]

        entity_text = _format_entities(data.get("entities"))
        if entity_text:
            sections.append(("Key People & Organizations", entity_text))

        sections.append(("Transcript", data["transcript"]))

        for title, content in sections:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(40, 10, title)
            pdf.ln(8)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, content or "")
            pdf.ln(10)

        output = BytesIO()
        pdf.output(output)
        output.seek(0)
        return output
